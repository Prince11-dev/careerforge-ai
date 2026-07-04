"""PDF and DOCX export service."""
import io
from typing import List, Dict
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

class DocumentGeneratorService:
    """Service for generating PDF and DOCX resumes."""

    def generate_pdf(self, sections: List[Dict], title: str = "Resume") -> bytes:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=0.6*inch,
            leftMargin=0.6*inch,
            topMargin=0.6*inch,
            bottomMargin=0.6*inch
        )

        styles = getSampleStyleSheet()

        # Custom styles
        title_style = ParagraphStyle(
            'ResumeTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=6,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )

        section_style = ParagraphStyle(
            'SectionTitle',
            parent=styles['Heading2'],
            fontSize=12,
            spaceAfter=4,
            spaceBefore=12,
            fontName='Helvetica-Bold',
            borderWidth=1,
            borderColor='black',
            borderPadding=2,
            backColor='lightgrey'
        )

        body_style = ParagraphStyle(
            'BodyText',
            parent=styles['BodyText'],
            fontSize=10,
            leading=14,
            spaceAfter=6
        )

        story = []

        for section in sorted(sections, key=lambda x: x.get("display_order", 0)):
            section_type = section.get("section_type", "")
            content = section.get("content", "")

            if not content.strip():
                continue

            # Add section title
            section_titles = {
                "header": "",
                "summary": "PROFESSIONAL SUMMARY",
                "skills": "TECHNICAL SKILLS",
                "experience": "PROFESSIONAL EXPERIENCE",
                "projects": "PROJECTS",
                "education": "EDUCATION",
                "certifications": "CERTIFICATIONS"
            }

            if section_type != "header" and section_titles.get(section_type):
                story.append(Paragraph(section_titles[section_type], section_style))
                story.append(Spacer(1, 4))

            # Format content
            for line in content.split("\n"):
                line = line.strip()
                if not line:
                    story.append(Spacer(1, 4))
                    continue

                if section_type == "header":
                    if line == content.split("\n")[0]:
                        story.append(Paragraph(line, title_style))
                    else:
                        story.append(Paragraph(line, ParagraphStyle('Contact', parent=body_style, alignment=TA_CENTER)))
                else:
                    story.append(Paragraph(line, body_style))

            story.append(Spacer(1, 8))

        doc.build(story)
        buffer.seek(0)
        return buffer.read()

    def generate_docx(self, sections: List[Dict], title: str = "Resume") -> bytes:
        doc = Document()

        # Set narrow margins
        sections_doc = doc.sections[0]
        sections_doc.top_margin = Inches(0.5)
        sections_doc.bottom_margin = Inches(0.5)
        sections_doc.left_margin = Inches(0.6)
        sections_doc.right_margin = Inches(0.6)

        for section in sorted(sections, key=lambda x: x.get("display_order", 0)):
            section_type = section.get("section_type", "")
            content = section.get("content", "")

            if not content.strip():
                continue

            section_titles = {
                "header": "",
                "summary": "PROFESSIONAL SUMMARY",
                "skills": "TECHNICAL SKILLS",
                "experience": "PROFESSIONAL EXPERIENCE",
                "projects": "PROJECTS",
                "education": "EDUCATION",
                "certifications": "CERTIFICATIONS"
            }

            if section_type != "header" and section_titles.get(section_type):
                heading = doc.add_heading(section_titles[section_type], level=2)
                heading.paragraph_format.space_after = Pt(4)
                heading.paragraph_format.space_before = Pt(12)

            for line in content.split("\n"):
                line = line.strip()
                if not line:
                    continue

                if section_type == "header":
                    if line == content.split("\n")[0]:
                        p = doc.add_paragraph()
                        run = p.add_run(line)
                        run.bold = True
                        run.font.size = Pt(16)
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        p = doc.add_paragraph(line)
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p.paragraph_format.space_after = Pt(2)
                else:
                    p = doc.add_paragraph(line)
                    p.paragraph_format.space_after = Pt(3)

            if section_type != "header":
                doc.add_paragraph()  # Add spacing between sections

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

document_generator_service = DocumentGeneratorService()
